# -*- coding: utf-8 -*-
"""Generate a complete Word document for the Technocas CRM project."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x6D, 0x28, 0xD9)      # violet
DARK = RGBColor(0x1E, 0x29, 0x3B)
GREY = RGBColor(0x64, 0x74, 0x8B)

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = BRAND if level == 1 else DARK
    return p

def para(text, bold=False, italic=False, color=None, size=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    if color: run.font.color.rgb = color
    if size: run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        _shade(hdr[i], '6D28D9')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x0B,0x3D,0x2E)
    _shade_para(p, 'F1F5F9')
    p.paragraph_format.left_indent = Inches(0.15)
    return p

def _shade_para(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    pPr.append(shd)

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Technocas CRM')
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = BRAND
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Complete Project & Database Documentation')
r.font.size = Pt(14); r.font.color.rgb = GREY
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('AI-assisted CRM for a custom apparel print shop (Decoinks)\nLast updated: 2026-06-12')
r.font.size = Pt(10); r.font.color.rgb = GREY
doc.add_paragraph()

# ============================================================
# 1. OVERVIEW
# ============================================================
h('1. What this project is', 1)
para('A single-inbox, real-time, AI-assisted CRM for a custom apparel print shop '
     '(Decoinks — hoodies, t-shirts, jerseys, DTF transfers, embroidery).')
bullet('Pulls Facebook Messenger + Instagram DMs into ONE inbox.')
bullet('Every conversation is auto-converted into a lead (idempotent).')
bullet('Agents get an AI Supervisor beside each chat: intent, sentiment, recommended reply, translation, summary, agent scoring.')
bullet('All data lives in a remote PostgreSQL database; every message is also embedded into a Qdrant vector database for semantic search / AI memory.')

# ============================================================
# 2. TECH STACK
# ============================================================
h('2. Technology stack', 1)
table(['Layer', 'Technology'], [
    ['Frontend', 'React 18 + Vite 5 + React Router 6 + TailwindCSS 3'],
    ['Backend', 'Node.js (ESM) + Express'],
    ['Primary database', 'PostgreSQL (pg driver)'],
    ['Vector database', 'Qdrant (REST API)'],
    ['AI', 'OpenAI — gpt-4o-mini (chat) + text-embedding-3-small (1536-dim embeddings)'],
    ['Realtime', 'SSE (/api/stream) + polling (conversations 4s, Meta pull 10s)'],
    ['Auth', 'JWT (bcrypt password hashing)'],
    ['Deploy', 'Docker Compose (nginx frontend + Node backend)'],
], widths=[1.7, 4.5])

# ============================================================
# 3. ARCHITECTURE
# ============================================================
h('3. Architecture (high level)', 1)
para('Facebook Messenger / Instagram DM  →  Express backend (Node)  →  PostgreSQL + Qdrant  →  React frontend.', italic=True)
para('Data store pattern (server/db.js): On boot the backend connects to Postgres, ensures the schema, '
     'and loads every row into an in-memory cache. Reads are served from memory (instant, synchronous). '
     'Writes update memory immediately AND write through to Postgres on an independent, non-blocking queue, '
     'so a slow remote write never blocks the UI. Every message is additionally embedded into Qdrant.')

# ============================================================
# 4. DATABASES — FULL DETAIL
# ============================================================
h('4. Databases — complete reference', 1)

h('4.1  PostgreSQL (primary database)', 2)
table(['Property', 'Value'], [
    ['Type', 'PostgreSQL (Docker container "decoinks-postgres")'],
    ['Host / IP', '31.97.110.197'],
    ['Port', '5436'],
    ['Database name', 'decoinks_db'],
    ['User', 'decoinks'],
    ['Connection string', 'server/.env  →  DATABASE_URL (gitignored, never committed)'],
    ['Pool config', 'max 12 clients · 5s connect timeout · 8s query/statement timeout · 30s idle'],
    ['Driver', 'node-postgres (pg ^8.21)'],
], widths=[1.9, 4.3])

para('Schema design (Phase 1 — pragmatic):', bold=True, space_after=2)
para('One table per resource. Every row has the SAME four columns; all real fields live inside the JSONB "doc".')
code('id  TEXT  PRIMARY KEY\ndoc         JSONB  NOT NULL\ncreated_at  TIMESTAMPTZ  DEFAULT now()\nupdated_at  TIMESTAMPTZ  DEFAULT now()')
para('Two extra key/value tables: settings(key, value) and meta_kv(key, value). '
     'JSONB values are written with JSON.stringify + ::jsonb so string tokens stay valid JSON.')

para('Tables and the exact fields stored in each row\'s "doc":', bold=True, space_after=2)
table(['Table', 'Rows', 'Fields stored inside doc (JSONB)'], [
    ['users', '2', 'id, name, email, role, password_hash, created_at'],
    ['customers', '12', 'id, name, company, email, phone, channel, tier, type, spend, orders, health, health_label, owner, loc, role, last_order, activity_days, activity_ago, avatar, initials, created_at'],
    ['conversations', '113', 'id, name, phone, company, channel, source, status, status_bg, status_icon, unread, bookmarked, assigned_to, customer_id, last_ts, list_preview, list_time, avatar_bg, channel_bg, initials, profile_pic, meta_recipient_id, created_at  +  AI cache: summary, summary_count, summary_at'],
    ['messages', '1899', 'id, conversation_id, dir (in/out/note), text, via, time, agent, created_at'],
    ['leads', '113', 'id, name, company, agent, status, statusCls, score, units, value, product, pipeline, pipelineCls, badge, source, source_type, conversation_id, av, initials, created, createdTime, created_at'],
    ['notes', '3', 'id, customer_id, title, body, category, author, pinned, date, created_at'],
    ['receipts', '8', 'id, receipt_no, order_no, customer, customer_orders, amount, method, method_icon, status, date, time, note, note2, created_at'],
    ['artworks', '8', 'id, name, type, product, order_no, customers, date, fav, bg, created_at'],
    ['orders', '0', '(defined, currently empty) order_no, products, ...'],
    ['payments', '0', '(defined, currently empty) invoice_no, order_no, description, ...'],
    ['webhook_events', '0+', 'id, source, received_at, event (raw Meta webhook payloads)'],
    ['settings', 'kv', 'key app_settings  →  { business{...}, notifications{...} }'],
    ['meta_kv', 'kv', '_autoinc counters; Meta tokens / page info'],
], widths=[1.3, 0.6, 4.3])
para('Row counts are as of the last database check.', italic=True, color=GREY, size=9)

h('4.2  Qdrant (vector database)', 2)
table(['Property', 'Value'], [
    ['Host / IP', '31.97.110.197'],
    ['Port', '6333'],
    ['Connection', 'server/.env  →  QDRANT_URL, QDRANT_API_KEY'],
    ['Embeddings', 'OpenAI text-embedding-3-small (1536 dimensions, cosine distance)'],
], widths=[1.6, 4.6])
table(['Collection', 'Vector', 'Purpose', 'Payload fields'], [
    ['crm_messages', '1536 / cosine', 'Every chat message embedded → semantic search / AI memory. Auto-ingested on each insert; 1757 existing messages backfilled.', 'message_id, conversation_id, dir, via, time, text, created_at'],
    ['documents', '1536 / cosine', 'Knowledge-base snippets used for RAG by AI analysis & recommended replies.', 'text, title, category, language, author, access_level, doc_id'],
], widths=[1.1, 0.9, 2.4, 1.8])
para('Message point IDs = deterministic md5→UUID of the message id, so re-ingesting updates the point instead of duplicating it.')

h('4.3  How writes reach the database', 2)
bullet('Read: getAll / findById → served from the in-memory cache (synchronous, instant).')
bullet('Write: insert / update / remove / setSetting → update cache immediately, then enqueue an independent write-through to Postgres (no-op writes skipped to keep the 10s Meta poller cheap).')
bullet('flush() is timeout-bounded (600ms) so the UI never hangs on a slow remote write; SIGTERM/SIGINT flush pending writes before exit.')
bullet('Every saved message is also fire-and-forget embedded into Qdrant crm_messages.')

# ============================================================
# 5. API
# ============================================================
h('5. Backend API (server/index.js, port 3001)', 1)
para('All /api/* routes except login, health and webhooks require a JWT (Authorization: Bearer <token>).')

para('Auth & users', bold=True, space_after=2)
table(['Method', 'Path', 'Purpose'], [
    ['POST', '/api/auth/login', 'Email + password → JWT'],
    ['GET', '/api/auth/me', 'Current user'],
    ['PATCH', '/api/auth/me', 'Update own name / role'],
    ['POST', '/api/auth/password', 'Change password'],
    ['GET / POST', '/api/users', 'List / create team members'],
    ['PATCH / DELETE', '/api/users/:id', 'Edit / remove (cannot delete self)'],
], widths=[1.2, 2.0, 3.0])

para('Generic CRUD (auto-generated per resource)', bold=True, space_after=2)
para('GET /api/<r> (list + ?search= + field filters) · GET /api/<r>/:id · POST /api/<r> · PATCH /api/<r>/:id · DELETE /api/<r>/:id  —  for: customers, leads, notes, orders, payments, receipts, artworks, conversations.')

para('Conversations, messages & stats', bold=True, space_after=2)
table(['Method', 'Path', 'Purpose'], [
    ['GET', '/api/conversations/:id/messages', 'Messages of a conversation (chat order)'],
    ['POST', '/api/conversations/:id/messages', 'Add a message / note (also → Qdrant)'],
    ['POST', '/api/conversations/:id/read', 'Mark read (clear unread)'],
    ['GET', '/api/stats/dashboard', 'Dashboard counters'],
], widths=[1.0, 2.7, 2.5])

para('Meta (Facebook / Instagram)', bold=True, space_after=2)
para('GET /api/meta/status · POST /api/meta/connect · /connect-app · /disconnect · /send · /lookup · /sync (manual pull) · GET|POST /api/webhooks/meta (verify + receive). '
     'A background poller pulls Messenger/Instagram conversations every 10 seconds and auto-captures leads.')

para('ManyChat (legacy / optional)', bold=True, space_after=2)
para('/api/manychat/status | connect | disconnect | page | send | lookup, /api/manychat/subscribers/:id, /api/webhooks/manychat.')

para('Leads', bold=True, space_after=2)
para('POST /api/leads/backfill — convert all existing conversations → leads.')

para('AI Supervisor & vectors', bold=True, space_after=2)
table(['Method', 'Path', 'Purpose'], [
    ['GET', '/api/ai/status', 'Is OpenAI / Qdrant configured'],
    ['GET', '/api/ai/analyze/:id', 'Full analysis (intent, sentiment, insights, scores, objections, lead prediction, recommended reply) + RAG'],
    ['POST', '/api/ai/recommend-reply/:id', 'Reply to UNANSWERED customer messages (after agent\'s last reply)'],
    ['GET', '/api/ai/summary/:id', 'Return cached conversation summary (+ new-message count)'],
    ['POST', '/api/ai/summary/:id', 'Generate (full) or incrementally update + save summary'],
    ['POST', '/api/ai/translate-assist', 'Explain last message in simple English + suggest reply (EN + native)'],
    ['POST', '/api/translate', 'Translate text (OpenAI, MyMemory fallback)'],
    ['POST', '/api/ai/ingest', 'Add a Knowledge-Base doc → Qdrant documents'],
    ['POST', '/api/ai/ingest-messages', 'Backfill all messages → Qdrant crm_messages'],
    ['GET', '/api/qdrant/status', 'Qdrant health + collections'],
], widths=[0.9, 2.3, 3.0])

para('Settings & misc', bold=True, space_after=2)
para('GET|PUT /api/settings (business + notifications) · GET /api/stream (SSE realtime) · GET /api/health.')

# ============================================================
# 6. PAGES
# ============================================================
h('6. Frontend pages (routes)', 1)
table(['Route', 'Page', 'What it does'], [
    ['/', 'Login', 'Email + password → JWT'],
    ['/dashboard', 'Dashboard', 'MAIN 3-panel inbox: filters · chat · AI Supervisor'],
    ['/inbox', '(redirect)', '→ /dashboard'],
    ['/customers', 'Customers', 'Customer list (search / filter)'],
    ['/customer-360', 'Customer 360', 'Full customer profile (hidden from sidebar)'],
    ['/leads', 'Leads', 'Lead pipeline list'],
    ['/leads/:id', 'Lead Details', 'Single lead'],
    ['/orders', 'Orders', 'Orders list'],
    ['/receipts', 'Receipts', 'Receipts list'],
    ['/follow-ups', 'Follow-Ups', 'Follow-up tasks'],
    ['/campaigns', 'Campaigns', 'Campaigns'],
    ['/reports', 'Reports', 'Reporting'],
    ['/settings', 'Settings', 'Business + notification settings'],
    ['/team', 'Team', 'Team CRUD (create login accounts, roles)'],
    ['/artwork-vault', 'Artwork Vault', 'Saved artwork'],
    ['/connect-meta', 'Meta Connect', 'Connect Facebook / Instagram'],
    ['/integrations', 'Integrations', 'Integration setup'],
], widths=[1.4, 1.3, 3.5])
para('All routes except / are protected by <RequireAuth>.', italic=True, color=GREY, size=9)

para('The Dashboard — three resizable panels:', bold=True, space_after=2)
bullet('1) Filters / conversation list — search, "All Inboxes" channel dropdown, Filters form (Channel, Agent, Lead Status, Tags, Date From/To), Latest/Oldest sort, view tabs (All / Unassigned / Mentions / Bookmarks). All filter the live backend list.')
bullet('2) Conversation — chat thread; send replies (Meta convs routed to Graph API), notes, status, bookmark, assign-to-me.')
bullet('3) AI Supervisor tabs: Responses (AI Recommended Reply for unanswered messages) · Translation · Summary (saved, incremental) · Actions (agent score, interpretation) · Designer Jobs · Intent & Insights.')

# ============================================================
# 7. FILE MAP
# ============================================================
h('7. Backend file map (server/)', 1)
table(['File', 'Role'], [
    ['index.js', 'All Express routes, Meta polling, AI endpoints, lead capture, message→Qdrant ingest'],
    ['db.js', 'PostgreSQL store (in-memory cache + write-through), synchronous API'],
    ['ai.js', 'OpenAI helper — embed(), chatJSON(), chatText()'],
    ['qdrant.js', 'Qdrant REST client — ensureCollection, upsert, search, payload index'],
    ['meta.js', 'Meta Graph API client (conversations, send, profiles, token exchange)'],
    ['manychat.js', 'ManyChat client (legacy)'],
    ['migrate-to-pg.js', 'One-time data.json → Postgres migration'],
    ['seed.js / qdrant-setup.js', 'Seeding / Qdrant bootstrap'],
], widths=[1.8, 4.4])

# ============================================================
# 8. AUTH + ENV + RUN + DEPLOY
# ============================================================
h('8. Authentication', 1)
bullet('Login → bcrypt-verified → JWT, stored client-side, sent as Authorization: Bearer.')
bullet('Per-user roles (Team page manages accounts). A user cannot delete themselves.')
bullet('Default login: info@technocas.com.')

h('9. Environment variables (server/.env — gitignored)', 1)
code('DATABASE_URL=postgresql://decoinks:<password>@31.97.110.197:5436/decoinks_db\n'
     'QDRANT_URL=http://31.97.110.197:6333\n'
     'QDRANT_API_KEY=<key>\n'
     'OPENAI_API_KEY=<key>\n'
     '# optional: OPENAI_CHAT_MODEL, OPENAI_EMBED_MODEL, PORT, JWT secret')
para('Secrets are never committed — .env, server/.env and server/data.json* are gitignored.')

h('10. Running locally', 1)
code('npm install            # frontend deps\n'
     'npm run server:install # backend deps\n'
     'npm run dev:all        # backend (3001) + Vite frontend (5173)')
para('Open http://localhost:5173 · API at http://localhost:3001. '
     'Backend scripts: dev (watch), start, seed, migrate, qdrant:setup.')

h('11. Deployment (Docker, server 31.97.110.197)', 1)
bullet('frontend — Vite build served by nginx (port 8190:80), proxies /api → backend.')
bullet('backend — Node, reads server/.env, reaches DB/Qdrant via host.docker.internal.')
code('cd ~/technocas-crm\ngit pull\ndocker compose up -d --build\ndocker compose logs -f backend')
para('Open http://31.97.110.197:8190. Repo: https://github.com/RahulAi2004/Technoas-Crm')

# ============================================================
# 12. DATA FLOWS
# ============================================================
h('12. Core data flows', 1)
para('Incoming message → lead → AI', bold=True, space_after=2)
bullet('Poller (10s) or webhook receives a Messenger/Instagram message.')
bullet('Conversation is upserted; a lead is auto-created (idempotent by conversation id).')
bullet('Message saved to Postgres AND embedded into Qdrant crm_messages.')
bullet('Frontend polls (4s) and shows it; AI Supervisor can analyze, summarize, or draft a reply.')
para('AI Recommended Reply', bold=True, space_after=2)
para('Looks only at the customer messages AFTER the agent\'s last reply (the unanswered ones), uses the full chat '
     '+ Knowledge Base as context, replies in the customer\'s language. No AI cost when the agent already replied to the latest message.')
para('Conversation Summary', bold=True, space_after=2)
para('Generated on demand, persisted on the conversation (summary, summary_count, summary_at). '
     'Re-runs are incremental: only the new messages since the last summary are sent to the model and merged in.')

# ============================================================
# 13. ROADMAP + SECURITY
# ============================================================
h('13. The bigger vision (roadmap)', 1)
para('Current state is Phase 1: app moved from a JSON file to remote PostgreSQL + Qdrant, with AI features live.')
para('Documented long-term architecture: Channels → Chatwoot → n8n → OpenAI → PostgreSQL + pgvector → Qdrant (RAG) → NestJS → React, '
     'with ~36 normalized tables, per-customer AI memory and 3-tier support. '
     'Phase 2 = normalize the JSONB doc tables into typed columns; later phases add pgvector, n8n automation and deeper AI/RAG.')

h('14. Security notes', 1)
bullet('Postgres/Qdrant ports are currently exposed directly on the public server IP — plan to move the app onto the server or use an SSH tunnel / firewall.')
bullet('All API secrets live only in server/.env (gitignored). Never print tokens in screenshots or logs.')
bullet('JWT + bcrypt for auth; passwords stored only as password_hash.')

doc.save('Technocas-CRM-Documentation.docx')
print('saved Technocas-CRM-Documentation.docx')
