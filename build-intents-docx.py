# -*- coding: utf-8 -*-
"""Build a Word doc of chatbot intents (grouped by category) from intents.json."""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRAND = RGBColor(0x6D, 0x28, 0xD9)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
GREY  = RGBColor(0x64, 0x74, 0x8B)

data = json.load(open('intents.json', encoding='utf-8'))
intents = data['intents']
src = data.get('source', {})

cats = {}
for it in intents:
    cats.setdefault(it.get('category', 'General'), []).append(it)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Decoinks — Chatbot Intents'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BRAND
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run(f"{len(intents)} distinct intents grouped from {src.get('conversations','?')} conversations "
              f"({src.get('customerMessages','?')} customer messages, {src.get('rawQuestions','?')} raw questions)")
r.font.size = Pt(10); r.font.color.rgb = GREY
doc.add_paragraph()

h = doc.add_heading('Categories', level=1)
for run in h.runs: run.font.color.rgb = BRAND
for cat, items in cats.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{cat} ').bold = True
    p.add_run(f'({len(items)} intents)').font.color.rgb = GREY
doc.add_page_break()

n = 0
for cat, items in cats.items():
    h = doc.add_heading(cat, level=1)
    for run in h.runs: run.font.color.rgb = BRAND
    for it in items:
        n += 1
        pq = doc.add_paragraph()
        rq = pq.add_run(f'{n}. {it.get("question","")}')
        rq.bold = True; rq.font.size = Pt(11); rq.font.color.rgb = DARK
        pq.paragraph_format.space_before = Pt(6); pq.paragraph_format.space_after = Pt(1)
        # intent id
        pi = doc.add_paragraph(); ri = pi.add_run(f'intent: {it.get("intent","")}')
        ri.font.size = Pt(8.5); ri.font.color.rgb = BRAND; ri.font.name = 'Consolas'
        pi.paragraph_format.space_after = Pt(1)
        # examples
        ex = it.get('examples') or []
        if ex:
            pe = doc.add_paragraph()
            pe.add_run('e.g. ').italic = True
            re = pe.add_run(' · '.join(ex)); re.italic = True; re.font.size = Pt(9.5); re.font.color.rgb = GREY
            pe.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()

doc.save('Decoinks-Chatbot-Intents.docx')
print(f'saved Decoinks-Chatbot-Intents.docx ({len(intents)} intents in {len(cats)} categories)')
