# -*- coding: utf-8 -*-
"""Build a Word document of FAQs + best agent replies from faq.json."""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x6D, 0x28, 0xD9)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
GREY  = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x05, 0x96, 0x69)

data = json.load(open('faq.json', encoding='utf-8'))
faqs = data['faqs']
gen = data.get('generatedFrom', {})

# group by category, preserve first-seen order
cats = {}
for f in faqs:
    cats.setdefault(f.get('category', 'General'), []).append(f)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def shade(p, hexc):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc)
    pPr.append(shd)

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Decoinks — Customer FAQ & Agent Reply Guide'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BRAND
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run(f"{len(faqs)} unique questions extracted from {gen.get('conversations','?')} real conversations "
              f"({gen.get('messages','?')} customer messages) + AI-generated best replies")
r.font.size = Pt(10); r.font.color.rgb = GREY
doc.add_paragraph()

# Contents
h = doc.add_heading('Categories', level=1)
for run in h.runs: run.font.color.rgb = BRAND
for cat, items in cats.items():
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{cat} ').bold = True
    p.add_run(f'({len(items)})').font.color.rgb = GREY
doc.add_page_break()

n = 0
for cat, items in cats.items():
    h = doc.add_heading(cat, level=1)
    for run in h.runs: run.font.color.rgb = BRAND
    for f in items:
        n += 1
        # Question
        pq = doc.add_paragraph()
        rq = pq.add_run(f'Q{n}. {f["question"]}')
        rq.bold = True; rq.font.size = Pt(11); rq.font.color.rgb = DARK
        pq.paragraph_format.space_before = Pt(8); pq.paragraph_format.space_after = Pt(2)
        # Reply label
        pl = doc.add_paragraph(); rl = pl.add_run('Suggested reply:')
        rl.bold = True; rl.font.size = Pt(8.5); rl.font.color.rgb = GREEN
        pl.paragraph_format.space_after = Pt(0)
        # Reply body (shaded)
        pa = doc.add_paragraph(); ra = pa.add_run(f['reply']); ra.font.size = Pt(10)
        shade(pa, 'F1F5F9')
        pa.paragraph_format.space_after = Pt(6)
        pa.paragraph_format.left_indent = Pt(6); pa.paragraph_format.right_indent = Pt(6)
    doc.add_paragraph()

note = doc.add_paragraph()
r = note.add_run('Note: Replies are AI-generated from your real chat history. Fill any [bracketed] placeholders '
                 '(exact prices, turnaround days, links) with your actual values before using as canned responses.')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

doc.save('Decoinks-FAQ-Agent-Replies.docx')
print(f'saved Decoinks-FAQ-Agent-Replies.docx ({len(faqs)} FAQs in {len(cats)} categories)')
