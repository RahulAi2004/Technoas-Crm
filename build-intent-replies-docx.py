# -*- coding: utf-8 -*-
"""Word doc: each intent's question + real agent reply (from chats) + AI recommended reply."""
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x6D, 0x28, 0xD9)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
GREY  = RGBColor(0x64, 0x74, 0x8B)
GREEN = RGBColor(0x05, 0x96, 0x69)
BLUE  = RGBColor(0x1D, 0x4E, 0xD8)

data = json.load(open('intents-replies.json', encoding='utf-8'))
intents = data['intents']

cats = {}
for it in intents:
    cats.setdefault(it.get('category', 'General'), []).append(it)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

def shade(p, hexc):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexc)
    pPr.append(shd)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('Decoinks — Chatbot Intents & Replies'); r.bold = True; r.font.size = Pt(21); r.font.color.rgb = BRAND
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run(f"{data['total']} intents · {data.get('withRealReply',0)} with a real agent reply from chats. "
              "Each: the actual reply your team gave + an AI-recommended reply.")
r.font.size = Pt(10); r.font.color.rgb = GREY
doc.add_paragraph()

h = doc.add_heading('Categories', level=1)
for run in h.runs: run.font.color.rgb = BRAND
for cat, items in cats.items():
    p = doc.add_paragraph(style='List Bullet'); p.add_run(f'{cat} ').bold = True; p.add_run(f'({len(items)})').font.color.rgb = GREY
doc.add_page_break()

n = 0
for cat, items in cats.items():
    h = doc.add_heading(cat, level=1)
    for run in h.runs: run.font.color.rgb = BRAND
    for it in items:
        n += 1
        pq = doc.add_paragraph(); rq = pq.add_run(f'Q{n}. {it.get("question","")}')
        rq.bold = True; rq.font.size = Pt(11.5); rq.font.color.rgb = DARK
        pq.paragraph_format.space_before = Pt(9); pq.paragraph_format.space_after = Pt(2)

        # Reply 1: real agent reply
        l1 = doc.add_paragraph(); r1 = l1.add_run("1) Agent's actual reply (from chats):")
        r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = GREEN
        l1.paragraph_format.space_after = Pt(0)
        real = it.get('realReply') or '— (no direct example found in chats)'
        pr = doc.add_paragraph(); ra = pr.add_run(real); ra.font.size = Pt(10)
        if it.get('realReply'): shade(pr, 'ECFDF5')
        pr.paragraph_format.left_indent = Pt(6); pr.paragraph_format.right_indent = Pt(6); pr.paragraph_format.space_after = Pt(4)

        # Reply 2: AI recommended
        l2 = doc.add_paragraph(); r2 = l2.add_run('2) AI recommended reply:')
        r2.bold = True; r2.font.size = Pt(8.5); r2.font.color.rgb = BLUE
        l2.paragraph_format.space_after = Pt(0)
        pa = doc.add_paragraph(); aa = pa.add_run(it.get('aiReply') or '—'); aa.font.size = Pt(10)
        shade(pa, 'EFF6FF')
        pa.paragraph_format.left_indent = Pt(6); pa.paragraph_format.right_indent = Pt(6); pa.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

note = doc.add_paragraph()
r = note.add_run('Note: "Agent\'s actual reply" is pulled from your real conversation history (the reply that followed a '
                 'similar customer question). Fill any [bracketed] placeholders in AI replies before use.')
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = GREY

doc.save('Decoinks-Chatbot-Intents-Replies.docx')
print(f'saved Decoinks-Chatbot-Intents-Replies.docx ({data["total"]} intents)')
